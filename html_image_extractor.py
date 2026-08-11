"""
HTML Image Extractor - извлича вградените снимки от HTML и прави Word документ.

Един файл, който сам се грижи за всичко.

Библиотеките за Word - python-docx, lxml и Pillow - носят машинен код и
не могат да се препишат тук. Затова при първото пускане програмата ги
сваля сама в собствена папка и показва прозорче с хода. След това тръгва
веднага и работи без интернет. Нищо не се инсталира на ръка.

Какво прави:
  • намира вградените Base64 изображения и ги записва като PNG, изрязани
    точно както се показват в браузъра, с махнат плътен фон
  • пренася текста, списъците и таблиците в .docx
  • оформлението се чете от самия CSS на страницата, затова работи с
    произволен HTML, а не само с една страница

Документът излиза на бял фон, с шрифт Times New Roman 13.5 pt.

Пуска се с двойно щракване, а от командния ред приема и файлове:

    python3 html_image_extractor.py файл.html
"""

import base64
import binascii
import hashlib
import html
import importlib
import io
import json
import os
import platform
import queue
import re
import shutil
import subprocess
import sys
import tempfile
import threading
import urllib.error
import urllib.request
import zipfile
from pathlib import Path

import tkinter as tk
from tkinter import filedialog, messagebox, ttk


# =========================================================================
# Самоподготовка на библиотеките
# =========================================================================

# Word документът се прави с python-docx, lxml и Pillow. Те носят
# компилирани части и затова не могат да се препишат в този файл. Вместо
# това при първото пускане се свалят сами в отделна папка на потребителя.
# Следващите пускания са мигновени и минават без интернет.

APP_NAME = "HTML Image Extractor"

FOLDER_NAME = "HtmlImageExtractor"

# Модулът, който се проверява, и пакетът, който се сваля за него.
REQUIRED_MODULES = (
    ("lxml", "lxml"),
    ("PIL", "pillow"),
    ("docx", "python-docx"),
)

# Пакети, които вървят заедно с друг пакет.
COMPANION_PACKAGES = {
    "python-docx": ("typing-extensions",),
}

PYPI_INDEX = "https://pypi.org/simple/{}/"

PYPI_METADATA = "https://pypi.org/pypi/{}/json"

USER_AGENT = "html-image-extractor"

# име-версия[-построяване]-python-abi-платформа.whl
WHEEL_NAME = re.compile(
    r"^(?P<name>[^-]+)-(?P<version>[^-]+)"
    r"(?:-(?P<build>\d[^-]*))?"
    r"-(?P<python>[^-]+)-(?P<abi>[^-]+)-(?P<platform>[^-]+)\.whl$",
    re.IGNORECASE,
)


def needed_by_downloaded_packages():
    """
    Изброява частите от стандартната библиотека, които python-docx, lxml
    и Pillow ползват.

    Тази функция никога не се вика. Стои, защото при сглобяване с
    PyInstaller вътре влиза само това, което се вижда написано в кода.
    Пакетите, свалени чак след сглобяването, не могат да се обадят какво
    им трябва, затова нуждите им са изброени тук - иначе сглобеният файл
    ги сваля успешно, но после те не тръгват.
    """
    import __future__
    import argparse
    import array
    import atexit
    import bisect
    import calendar
    import cmath
    import collections.abc
    import contextlib
    import copy
    import csv
    import dataclasses
    import datetime
    import decimal
    import difflib
    import enum
    import fnmatch
    import fractions
    import functools
    import getpass
    import glob
    import gzip
    import heapq
    import hmac
    import inspect
    import itertools
    import logging
    import logging.handlers
    import math
    import mimetypes
    import numbers
    import operator
    import pathlib
    import pickle
    import posixpath
    import pprint
    import random
    import secrets
    import shlex
    import stat
    import string
    import struct
    import textwrap
    import time
    import traceback
    import types
    import typing
    import unicodedata
    import urllib.parse
    import uuid
    import warnings
    import weakref
    import xml.dom.minidom
    import xml.etree.ElementTree
    import xml.parsers.expat
    import xml.sax


def hidden_process_flags():
    """Пречи на Windows да мига с черен прозорец при външна команда."""
    if sys.platform.startswith("win"):
        return {"creationflags": getattr(subprocess, "CREATE_NO_WINDOW", 0)}

    return {}


def application_folder():
    """Папката, в която стои програмата - или .py файлът, или .exe файлът."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent

    return Path(__file__).resolve().parent


def library_key():
    """
    Отпечатък на средата.

    Свалените пакети съдържат машинен код и важат само за една версия на
    Python върху една архитектура, затова всяка комбинация си има папка.
    """
    return "py{}{}-{}-{}".format(
        sys.version_info[0],
        sys.version_info[1],
        sys.platform,
        (platform.machine() or "unknown").lower(),
    )


def library_folder():
    """Постоянното място за свалените библиотеки."""
    if sys.platform.startswith("win"):
        base = os.environ.get("LOCALAPPDATA") or os.path.expanduser("~")
    elif sys.platform == "darwin":
        base = os.path.join(
            os.path.expanduser("~"), "Library", "Application Support"
        )
    else:
        base = os.environ.get("XDG_DATA_HOME") or os.path.join(
            os.path.expanduser("~"), ".local", "share"
        )

    return Path(base) / FOLDER_NAME / "libs" / library_key()


def add_to_path(folder):
    text = str(folder)

    if text not in sys.path:
        sys.path.insert(0, text)


def missing_packages(reasons=None):
    """
    Връща пакетите, които в момента не могат да се заредят.

    В reasons се събира по коя причина всеки от тях не тръгва - иначе
    после остава само едно „липсва“ без обяснение.
    """
    importlib.invalidate_caches()

    missing = []

    for module_name, package_name in REQUIRED_MODULES:
        try:
            importlib.import_module(module_name)
        except Exception as error:
            missing.append(package_name)

            if reasons is not None:
                reasons[package_name] = "{}: {}".format(
                    type(error).__name__,
                    error,
                )

    return missing


# ---------------------------------------------------------------------------
# Избор на подходящ пакет за тази машина
# ---------------------------------------------------------------------------

def python_tags():
    major, minor = sys.version_info[:2]

    return {
        "cp{}{}".format(major, minor),
        "py{}{}".format(major, minor),
        "py{}".format(major),
        "cp{}".format(major),
    }


def abi_tags():
    major, minor = sys.version_info[:2]

    return {
        "none",
        "abi3",
        "cp{}{}".format(major, minor),
        "cp{}{}m".format(major, minor),
    }


def macos_version():
    try:
        pieces = platform.mac_ver()[0].split(".")

        return (int(pieces[0]), int(pieces[1]) if len(pieces) > 1 else 0)
    except (IndexError, ValueError):
        return (99, 0)


def platform_matches(tag):
    """Проверява дали един етикет за платформа върви на тази машина."""
    if tag == "any":
        return True

    machine = (platform.machine() or "").lower()

    if sys.platform.startswith("win"):
        if machine in ("amd64", "x86_64"):
            return tag == "win_amd64"

        if machine in ("arm64", "aarch64"):
            return tag == "win_arm64"

        return tag == "win32"

    if sys.platform == "darwin":
        match = re.match(r"^macosx_(\d+)_(\d+)_(.+)$", tag)

        if not match:
            return False

        needed = (int(match.group(1)), int(match.group(2)))

        if needed > macos_version():
            return False

        architecture = match.group(3)

        if machine == "arm64":
            return architecture in ("arm64", "universal2")

        return architecture in (
            "x86_64",
            "universal2",
            "intel",
            "fat64",
            "fat32",
            "universal",
        )

    if machine in ("x86_64", "amd64"):
        endings = ("_x86_64",)
    elif machine in ("aarch64", "arm64"):
        endings = ("_aarch64",)
    elif machine.startswith("armv7"):
        endings = ("_armv7l",)
    else:
        endings = ("_" + machine,)

    if not tag.startswith(("manylinux", "musllinux", "linux")):
        return False

    return tag.endswith(endings)


def python_allows(specifier):
    """Проверява requires-python, записан като '>=3.9, <4'."""
    if not specifier:
        return True

    current = sys.version_info[:3]

    for part in specifier.split(","):
        match = re.match(
            r"^\s*(==|!=|<=|>=|~=|<|>)\s*([0-9][0-9a-zA-Z.*+!-]*)\s*$",
            part,
        )

        if not match:
            continue

        operator = match.group(1)
        raw = match.group(2).split("+")[0]
        wildcard = raw.endswith(".*")

        if wildcard:
            raw = raw[:-2]

        wanted = tuple(
            int(piece) for piece in raw.split(".") if piece.isdigit()
        )

        if not wanted:
            continue

        here = current[: len(wanted)]

        while len(here) < len(wanted):
            here += (0,)

        if operator == ">=" and here < wanted:
            return False

        if operator == ">" and here <= wanted:
            return False

        if operator == "<=" and here > wanted:
            return False

        if operator == "<" and here >= wanted:
            return False

        if operator == "==" and here != wanted:
            return False

        if operator == "!=" and here == wanted:
            return False

        if operator == "~=" and here < wanted:
            return False

    return True


def tags_match(match):
    """Етикетите могат да са слети с точка: py2.py3-none-any"""
    if not set(match.group("python").split(".")) & python_tags():
        return False

    if not set(match.group("abi").split(".")) & abi_tags():
        return False

    return any(
        platform_matches(tag) for tag in match.group("platform").split(".")
    )


def version_key(version):
    """Подрежда версиите по числата, после по добавката след тях."""
    match = re.match(r"^(\d+(?:\.\d+)*)(.*)$", version)

    if not match:
        return ((0,), version)

    numbers = tuple(int(piece) for piece in match.group(1).split("."))

    return (numbers, match.group(2))


def is_stable(version):
    """Пробните издания - 12.5.0b1, 2.0rc1 - не стават за всеки ден."""
    match = re.match(r"^(\d+(?:\.\d+)*)(.*)$", version)

    if not match:
        return False

    rest = match.group(2)

    return rest == "" or rest.startswith(".post")


def choose_wheel(listing):
    """
    Избира най-новия пакет, който пасва на тази машина.

    Пробните издания се гледат само ако няма нито едно редовно - иначе
    една бета версия с по-голямо число би изместила готовата.
    """
    best = [None, None]

    for entry in listing:
        filename = entry.get("filename") or ""

        if not filename.lower().endswith(".whl") or entry.get("yanked"):
            continue

        match = WHEEL_NAME.match(filename)

        if not match or not tags_match(match):
            continue

        limit = entry.get("requires-python") or entry.get("requires_python")

        if not python_allows(limit):
            continue

        version = match.group("version")
        slot = 0 if is_stable(version) else 1
        key = version_key(version)

        if best[slot] is None or key > best[slot][0]:
            best[slot] = (key, entry, filename)

    chosen = best[0] or best[1]

    if chosen is None:
        return None

    return chosen[1], chosen[2]


# ---------------------------------------------------------------------------
# Сваляне
# ---------------------------------------------------------------------------

def open_url(url, accept=None):
    headers = {"User-Agent": USER_AGENT}

    if accept:
        headers["Accept"] = accept

    return urllib.request.urlopen(
        urllib.request.Request(url, headers=headers),
        timeout=60,
    )


def read_simple_index(package_name):
    """Официалният списък с файлове на пакета."""
    url = PYPI_INDEX.format(package_name.lower().replace("_", "-"))

    with open_url(url, "application/vnd.pypi.simple.v1+json") as response:
        data = json.loads(response.read().decode("utf-8"))

    return data.get("files") or []


def read_legacy_index(package_name):
    """Резервен списък, ако новият отговор не се разчете."""
    with open_url(PYPI_METADATA.format(package_name)) as response:
        data = json.loads(response.read().decode("utf-8"))

    files = []

    for entries in (data.get("releases") or {}).values():
        for entry in entries:
            files.append(
                {
                    "filename": entry.get("filename"),
                    "url": entry.get("url"),
                    "requires-python": entry.get("requires_python"),
                    "yanked": entry.get("yanked"),
                    "hashes": entry.get("digests") or {},
                }
            )

    return files


def read_index(package_name):
    try:
        return read_simple_index(package_name)
    except (urllib.error.URLError, ValueError, OSError):
        return read_legacy_index(package_name)


def download(url, report):
    pieces = []
    received = 0

    with open_url(url) as response:
        try:
            size = int(response.headers.get("Content-Length") or 0)
        except (TypeError, ValueError):
            size = 0

        while True:
            piece = response.read(65536)

            if not piece:
                break

            pieces.append(piece)
            received += len(piece)
            report(received, size)

    return b"".join(pieces)


def extract_wheel(raw, target):
    """Разопакова пакета направо в папката, както прави pip --target."""
    target.mkdir(parents=True, exist_ok=True)
    root = target.resolve()

    with zipfile.ZipFile(io.BytesIO(raw)) as archive:
        for item in archive.infolist():
            if item.filename.endswith("/"):
                continue

            parts = item.filename.split("/")

            # Съдържанието на .data папката се разгъва в корена.
            if len(parts) > 2 and parts[0].endswith(".data"):
                if parts[1] not in ("purelib", "platlib"):
                    continue

                parts = parts[2:]

            if not parts or ".." in parts:
                continue

            destination = target.joinpath(*parts)

            if root not in destination.resolve().parents:
                continue

            destination.parent.mkdir(parents=True, exist_ok=True)

            with archive.open(item) as source:
                with open(str(destination), "wb") as output:
                    shutil.copyfileobj(source, output)

            mode = (item.external_attr >> 16) & 0o777

            # Архивите от Windows нямат права, там режимът се пропуска.
            if mode:
                os.chmod(str(destination), mode)


def fetch_libraries(packages, target, report):
    """Сваля и разопакова пакетите. Връща текст с грешка или None."""
    total = len(packages)

    for index, package_name in enumerate(packages):
        base = index / float(total)
        step = 1.0 / total

        report("Търся {}...".format(package_name), base)

        try:
            listing = read_index(package_name)
        except Exception as error:
            return "Няма връзка с pypi.org: {}".format(error)

        chosen = choose_wheel(listing)

        if chosen is None:
            return (
                "За {} няма готов пакет за Python {}.{} на тази машина."
            ).format(package_name, sys.version_info[0], sys.version_info[1])

        entry, filename = chosen

        def progress(received, size, name=filename, start=base, span=step):
            if size:
                report(
                    "Свалям {}".format(name),
                    start + span * 0.9 * received / size,
                )
            else:
                report("Свалям {}".format(name), None)

        try:
            raw = download(entry.get("url"), progress)
        except Exception as error:
            return "Свалянето на {} не стана: {}".format(filename, error)

        digest = (entry.get("hashes") or {}).get("sha256")

        if digest and hashlib.sha256(raw).hexdigest() != digest:
            return "Файлът {} се получи повреден.".format(filename)

        report("Разопаковам {}".format(filename), base + step * 0.95)

        try:
            extract_wheel(raw, target)
        except Exception as error:
            return "Разопаковането на {} не стана: {}".format(filename, error)

    report("Готово.", 1.0)

    return None


def install_with_pip(packages, target):
    """Резервен път през pip. Връща текст с грешка или None при успех."""
    if getattr(sys, "frozen", False):
        return "В сглобения файл няма pip."

    command = [
        sys.executable,
        "-m",
        "pip",
        "install",
        "--no-input",
        "--disable-pip-version-check",
        "--upgrade",
        "--target",
        str(target),
    ] + list(packages)

    for attempt in range(2):
        try:
            finished = subprocess.run(
                command,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                timeout=1800,
                **hidden_process_flags()
            )
        except (OSError, subprocess.SubprocessError) as error:
            return "pip не тръгна: {}".format(error)

        if finished.returncode == 0:
            return None

        output = (finished.stdout or b"").decode("utf-8", "replace")

        if attempt == 0 and "No module named pip" in output:
            try:
                subprocess.run(
                    [sys.executable, "-m", "ensurepip", "--default-pip"],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    timeout=600,
                    **hidden_process_flags()
                )
            except (OSError, subprocess.SubprocessError):
                return "pip липсва и не може да се създаде."

            continue

        return "pip спря с грешка:\n{}".format(output.strip()[-400:])

    return "pip не успя."


# ---------------------------------------------------------------------------
# Показване на хода
# ---------------------------------------------------------------------------

class ConsoleProgress:
    """Изписва хода в конзолата - за пускане от командния ред."""

    def run(self, work):
        # Свалянето вика доклад на всяко парче, затова се изписва само
        # когато има какво ново да се каже.
        last = [None]

        def report(message, fraction):
            step = message if fraction is None else (
                message,
                int(fraction * 100),
            )

            if step == last[0]:
                return

            last[0] = step

            if fraction is None:
                print("  {}".format(message))
            else:
                print("  {} {:3.0f}%".format(message, fraction * 100))

            sys.stdout.flush()

        return work(report)


class WindowProgress:
    """Малко прозорче с лента, докато свалянето върви в отделна нишка."""

    def run(self, work):
        outcome = {}
        updates = queue.Queue()

        root = tk.Tk()
        root.title(APP_NAME)
        root.resizable(False, False)
        root.protocol("WM_DELETE_WINDOW", lambda: None)

        frame = ttk.Frame(root, padding=24)
        frame.pack(fill="both", expand=True)

        ttk.Label(
            frame,
            text="Подготвям нужните компоненти",
            font=("Segoe UI", 14, "bold"),
        ).pack(anchor="w")

        ttk.Label(
            frame,
            text=(
                "Това се случва само първия път и трае около минута.\n"
                "Нужен е интернет. После програмата тръгва веднага."
            ),
            justify="left",
        ).pack(anchor="w", pady=(6, 14))

        bar = ttk.Progressbar(frame, length=440, maximum=100)
        bar.pack(fill="x")

        status = ttk.Label(frame, text="Започвам...", width=58, anchor="w")
        status.pack(anchor="w", pady=(10, 0))

        def report(message, fraction):
            updates.put((message, fraction))

        def worker():
            try:
                outcome["value"] = work(report)
            except Exception as error:
                outcome["value"] = str(error)

            updates.put(None)

        def drain():
            finished = False

            try:
                while True:
                    update = updates.get_nowait()

                    if update is None:
                        finished = True
                        continue

                    message, fraction = update
                    status.configure(text=message)

                    if fraction is not None:
                        bar.configure(value=max(0.0, min(1.0, fraction)) * 100)
            except queue.Empty:
                pass

            if finished:
                root.destroy()
            else:
                root.after(80, drain)

        threading.Thread(target=worker, daemon=True).start()
        root.after(80, drain)
        root.mainloop()

        return outcome.get("value")


def prepare_libraries():
    """
    Осигурява библиотеките за Word.

    Връща None при успех или текст с грешка, ако нищо не се е получило.
    Ако пакетите вече са налице, не се пипа мрежата и не се вижда нищо.
    """
    portable = application_folder() / "_libs" / library_key()
    private = library_folder()

    for folder in (portable, private):
        if folder.is_dir():
            add_to_path(folder)

    missing = missing_packages()

    if not missing:
        return None

    wanted = []

    for package_name in missing:
        for companion in COMPANION_PACKAGES.get(package_name, ()):
            if companion not in wanted:
                wanted.append(companion)

        if package_name not in wanted:
            wanted.append(package_name)

    try:
        private.mkdir(parents=True, exist_ok=True)
    except OSError as error:
        return "Няма достъп до папката {}: {}".format(private, error)

    add_to_path(private)

    def work(report):
        error = fetch_libraries(wanted, private, report)

        if error is None:
            return None

        report("Опитвам по друг начин, с pip...", None)

        if install_with_pip(wanted, private) is None:
            return None

        return error

    if len(sys.argv) > 1:
        print("Липсват компоненти, свалям ги еднократно...")
        error = ConsoleProgress().run(work)
    else:
        try:
            error = WindowProgress().run(work)
        except tk.TclError:
            error = ConsoleProgress().run(work)

    if error is not None:
        return error

    reasons = {}
    still = missing_packages(reasons)

    if still:
        return "Свалянето мина, но {} не тръгва: {}".format(
            ", ".join(still),
            " | ".join(reasons.values()),
        )

    return None


BOOTSTRAP_ERROR = prepare_libraries()


# Word документът изисква външни библиотеки. Ако липсват, приложението
# продължава да работи, но само с извличането на снимки.
try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt, RGBColor
    from lxml import html as lxml_html
    from PIL import Image, ImageChops, ImageDraw

    CONVERTER_ERROR = None
except ImportError as error:
    CONVERTER_ERROR = BOOTSTRAP_ERROR or str(error)


# =========================================================================
# Модел на CSS
# =========================================================================

# Свойства, които се предават на децата.
INHERITED_PROPERTIES = (
    "color",
    "text-align",
    "font-weight",
    "font-style",
    "text-decoration",
    "font-family",
    "list-style-type",
    "visibility",
)

# Селектори с тези части не се поддържат и правилото се пропуска.
UNSUPPORTED = re.compile(
    r"[>+~]"
    r"|::(?!marker)"
    r"|(?<!:):(?!:)(?!root\b|host\b)[a-z-]+"
)

COMMENT = re.compile(r"/\*.*?\*/", re.DOTALL)

TOKEN = re.compile(
    r"""
      \.(?P<cls>(?:\\.|[^\s.\#:\[])+)
    | \#(?P<id>(?:\\.|[^\s.\#:\[])+)
    | \[(?P<attr>[^\]]*)\]
    | ::(?P<pseudo>[a-zA-Z-]+)
    | :(?P<pseudo_class>[a-zA-Z-]+)
    | (?P<tag>\*|[a-zA-Z][\w-]*)
    | (?P<space>\s+)
    """,
    re.VERBOSE,
)


def unescape(text):
    """Tailwind изписва класовете екранирано: text-\\[\\#FFD100\\]"""
    return re.sub(r"\\(.)", r"\1", text)


def split_top_level(text, separator=","):
    """Разделя по запетая, но не вътре в скоби."""
    parts = []
    depth = 0
    current = []

    for character in text:
        if character in "([":
            depth += 1
        elif character in ")]":
            depth -= 1

        if character == separator and depth <= 0:
            parts.append("".join(current))
            current = []
        else:
            current.append(character)

    parts.append("".join(current))

    return [part.strip() for part in parts if part.strip()]


class Component:
    """Една част от селектора, например div.flex#main"""

    __slots__ = ("tag", "identifier", "classes", "attributes", "pseudo")

    def __init__(self):
        self.tag = None
        self.identifier = None
        self.classes = []
        self.attributes = []
        self.pseudo = None

    def matches(self, element):
        if self.tag and self.tag != "*" and element.tag != self.tag:
            return False

        if self.identifier and element.get("id") != self.identifier:
            return False

        if self.classes:
            present = set((element.get("class") or "").split())

            if not present.issuperset(self.classes):
                return False

        for name, value in self.attributes:
            if element.get(name) is None:
                return False

            if value is not None and element.get(name) != value:
                return False

        return True

    @property
    def key(self):
        """Ключ за бързо търсене - най-характерната част на компонента."""
        if self.identifier:
            return ("id", self.identifier)

        if self.classes:
            return ("class", self.classes[0])

        if self.tag and self.tag != "*":
            return ("tag", self.tag)

        return ("all", None)


class Rule:
    __slots__ = ("components", "declarations", "specificity", "order", "pseudo")

    def __init__(self, components, declarations, order):
        self.components = components
        self.declarations = declarations
        self.order = order
        self.pseudo = components[-1].pseudo

        identifiers = sum(1 for part in components if part.identifier)
        classes = sum(
            len(part.classes) + len(part.attributes) for part in components
        )
        tags = sum(1 for part in components if part.tag and part.tag != "*")

        self.specificity = (identifiers, classes, tags)

    def matches(self, element):
        if not self.components[-1].matches(element):
            return False

        if len(self.components) == 1:
            return True

        remaining = list(self.components[:-1])
        node = element.getparent()

        while node is not None and remaining:
            if isinstance(node.tag, str) and remaining[-1].matches(node):
                remaining.pop()

            node = node.getparent()

        return not remaining


def parse_component(text):
    component = Component()
    position = 0

    while position < len(text):
        match = TOKEN.match(text, position)

        if match is None:
            return None

        position = match.end()

        if match.group("cls"):
            component.classes.append(unescape(match.group("cls")))
        elif match.group("id"):
            component.identifier = unescape(match.group("id"))
        elif match.group("tag"):
            component.tag = match.group("tag").lower()
        elif match.group("attr"):
            raw = match.group("attr")

            if "=" in raw:
                name, _, value = raw.partition("=")
                component.attributes.append(
                    (name.strip(), value.strip().strip("\"'"))
                )
            else:
                component.attributes.append((raw.strip(), None))
        elif match.group("pseudo"):
            component.pseudo = match.group("pseudo")
        elif match.group("pseudo_class"):
            # :root и :host важат за най-горния елемент
            component.pseudo = match.group("pseudo_class")

    return component


def parse_selector(text):
    """Връща списък от компоненти или None, ако селекторът не се поддържа."""
    if UNSUPPORTED.search(text):
        return None

    components = []

    for piece in text.split():
        component = parse_component(piece)

        if component is None:
            return None

        components.append(component)

    return components or None


def parse_declarations(text):
    declarations = {}

    for part in split_top_level(text, ";"):
        if ":" not in part:
            continue

        name, _, value = part.partition(":")
        name = name.strip().lower()
        value = value.strip()

        if value.endswith("!important"):
            value = value[: -len("!important")].strip()

        if name:
            declarations[name] = value

    return declarations


class Stylesheet:
    def __init__(self):
        self.rules = []
        self.index = {}
        self.variables = {}
        self.order = 0
        self._cache = {}

    # -- четене ----------------------------------------------------------

    def add(self, source):
        self.parse_block(COMMENT.sub("", source or ""))

    def parse_block(self, text):
        position = 0
        length = len(text)

        while position < length:
            opening = text.find("{", position)

            if opening == -1:
                break

            # Правила без блок, например @import "..."; или @layer a, b;
            statement = text.find(";", position)

            if statement != -1 and statement < opening:
                position = statement + 1
                continue

            selector = text[position:opening].strip()
            depth = 1
            cursor = opening + 1

            while cursor < length and depth:
                if text[cursor] == "{":
                    depth += 1
                elif text[cursor] == "}":
                    depth -= 1

                cursor += 1

            body = text[opening + 1: cursor - 1]
            position = cursor

            if selector.startswith("@"):
                name = selector.split(None, 1)[0].lower()

                # Условните правила зависят от екрана - взимаме основния вид.
                if name in ("@layer", "@scope", "@document"):
                    self.parse_block(body)

                continue

            self.add_rule(selector, body)

    def add_rule(self, selector, body):
        declarations = parse_declarations(body)

        if not declarations:
            return

        for single in split_top_level(selector):
            components = parse_selector(single)

            if components is None:
                continue

            last = components[-1]

            if last.pseudo in ("root", "host"):
                self.collect_variables(declarations)
                continue

            self.order += 1
            rule = Rule(components, declarations, self.order)

            self.index.setdefault(last.key, []).append(rule)
            self.rules.append(rule)

    def collect_variables(self, declarations):
        for name, value in declarations.items():
            if name.startswith("--"):
                self.variables[name] = value

    # -- прилагане -------------------------------------------------------

    def resolve(self, value, depth=0):
        """Заменя var(--име, резервно) със стойността от :root."""
        if not value or "var(" not in value or depth > 6:
            return value

        def replace(match):
            inner = match.group(1)
            name, _, fallback = inner.partition(",")
            name = name.strip()

            if name in self.variables:
                return self.variables[name]

            return fallback.strip()

        return self.resolve(re.sub(r"var\(([^()]*)\)", replace, value), depth + 1)

    def candidates(self, element):
        keys = [("all", None), ("tag", element.tag)]

        identifier = element.get("id")

        if identifier:
            keys.append(("id", identifier))

        for name in (element.get("class") or "").split():
            keys.append(("class", name))

        seen = []

        for key in keys:
            seen.extend(self.index.get(key, ()))

        return seen

    def declarations_for(self, element, pseudo=None):
        """Всички правила за елемента, слети по специфичност."""
        # Ключът пази самия елемент, а не id() - lxml раздава обвивките
        # наново и един и същи id() може да се падне на различни възли.
        cache_key = (element, pseudo)

        if cache_key in self._cache:
            return self._cache[cache_key]

        matched = [
            rule
            for rule in self.candidates(element)
            if rule.pseudo == pseudo and rule.matches(element)
        ]
        matched.sort(key=lambda rule: (rule.specificity, rule.order))

        declarations = {}

        for rule in matched:
            declarations.update(rule.declarations)

        if pseudo is None:
            declarations.update(parse_declarations(element.get("style", "")))

        resolved = {
            name: self.resolve(value) for name, value in declarations.items()
        }

        self._cache[cache_key] = resolved

        return resolved


# =========================================================================
# Преобразуване към Word
# =========================================================================

# ---------------------------------------------------------------------------
# Настройки на документа
# ---------------------------------------------------------------------------

BASE_FONT = "Times New Roman"
FONT_SIZE_PT = 13.5

PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
PAGE_MARGIN_INCHES = 0.8

CONTENT_WIDTH_INCHES = (PAGE_WIDTH_MM / 25.4) - (2 * PAGE_MARGIN_INCHES)

# Ширината на браузъра се приравнява на ширината на листа. Така всички
# съотношения от страницата се запазват, а съдържанието се побира.
REFERENCE_VIEWPORT_PX = 1200.0
INCHES_PER_CSS_PX = CONTENT_WIDTH_INCHES / REFERENCE_VIEWPORT_PX

# Страницата е бяла, затова светлите цветове се затъмняват до четими.
# Прагът е избран така, че наситени акценти като златното #F3B007 да минат
# непроменени, а бялото и светлосивото да станат черни.
MAX_TEXT_LUMINANCE = 185
NEUTRAL_RANGE = 40

# Плътният фон на изображенията се маха по свързаност с ръба, за да не
# пострадат тъмните детайли вътре в рисунката. Пикселите, отдалечени от
# фона по-малко от долния праг, са шум от компресията и падат изцяло -
# иначе оставят сива мъгла около рисунката.
BACKDROP_FLOOR = 14
BACKDROP_TOLERANCE = 60

# Иконите, рисувани за тъмен екран, са невидими върху бяла страница.
# Когато е включено, знакът им се пребоядисва в тъмно, а подложката пада.
FIX_INVISIBLE_IMAGES = True

# Таван на приноса на текста при пресмятане на ширините на колоните.
# Без него една дълга клетка би изяла целия ред.
TEXT_WIDTH_CAP = 700.0

LIST_TYPES = {
    "disc": "•",
    "circle": "◦",
    "square": "▪",
}

SKIP_TAGS = {"script", "style", "head", "meta", "link", "title", "noscript"}

BLOCK_TAGS = {
    "p", "h1", "h2", "h3", "h4", "h5", "h6",
    "ul", "ol", "li", "table", "svg", "hr",
}

HEADING_TAGS = ("h1", "h2", "h3", "h4", "h5", "h6")

DATA_IMAGE_PATTERN = re.compile(
    r"data:image/(?P<mime>[a-zA-Z0-9.+_-]+)(?:;[^,]*)?;base64,(?P<data>.+)",
    re.IGNORECASE | re.DOTALL,
)


# ---------------------------------------------------------------------------
# Цветове и мерки
# ---------------------------------------------------------------------------

def paragraph_alignment(name):
    """Подравняването от CSS към това на Word."""
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "start": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "end": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(name)


def parse_color(value):
    """#RRGGBB, #RGB, rgb(...) -> (r, g, b)"""
    if not value:
        return None

    value = value.strip()

    named = {
        "black": (0, 0, 0),
        "white": (255, 255, 255),
        "red": (255, 0, 0),
        "transparent": None,
    }

    if value.lower() in named:
        return named[value.lower()]

    match = re.fullmatch(r"#([0-9a-fA-F]{6})", value)

    if match:
        raw = match.group(1)
        return tuple(int(raw[index:index + 2], 16) for index in (0, 2, 4))

    match = re.fullmatch(r"#([0-9a-fA-F]{3})", value)

    if match:
        return tuple(int(character * 2, 16) for character in match.group(1))

    match = re.fullmatch(r"rgba?\(([^)]+)\)", value)

    if match:
        numbers = re.findall(r"\d+(?:\.\d+)?", match.group(1))

        if len(numbers) >= 3:
            return tuple(int(float(number)) for number in numbers[:3])

    return None


def luminance(color):
    red, green, blue = color[:3]
    return 0.299 * red + 0.587 * green + 0.114 * blue


def readable_on_white(color):
    """Затъмнява светлите цветове, за да се четат върху бяла страница."""
    if color is None:
        return (0, 0, 0)

    level = luminance(color)

    # Белите и сивите тонове нямат оттенък за пазене - стават черни.
    if max(color) - min(color) < NEUTRAL_RANGE and level > 120:
        return (0, 0, 0)

    if level <= MAX_TEXT_LUMINANCE:
        return color

    factor = MAX_TEXT_LUMINANCE / level

    return tuple(int(round(channel * factor)) for channel in color)


def to_hex(color):
    return "{:02X}{:02X}{:02X}".format(*color)


def parse_length(value, reference=None):
    """CSS дължина в пиксели. Поддържа px, rem, em, pt и проценти."""
    if not value:
        return None

    value = value.strip().lower()

    if value in ("auto", "inherit", "initial", "none"):
        return None

    match = re.fullmatch(r"(-?\d+(?:\.\d+)?)\s*(px|rem|em|pt|%)?", value)

    if not match:
        number = re.search(r"-?\d+(?:\.\d+)?", value)
        return float(number.group(0)) if number else None

    amount = float(match.group(1))
    unit = match.group(2) or "px"

    if unit == "px":
        return amount

    if unit == "rem":
        return amount * 16.0

    if unit == "em":
        return amount * (reference or 16.0)

    if unit == "pt":
        return amount * 4.0 / 3.0

    if unit == "%":
        return (reference * amount / 100.0) if reference else None

    return amount


def normalize_text(text):
    if not text:
        return ""

    return re.sub(r"\s+", " ", text.replace("\xa0", " "))


def visible_text(element):
    """
    Текстът на елемента без съдържанието на script и style.

    Вграденият в SVG стил иначе минава за текст и обърква сметките.
    """
    pieces = []

    def collect(node):
        if not isinstance(node.tag, str):
            return

        if node.tag in SKIP_TAGS:
            return

        if node.text:
            pieces.append(node.text)

        for child in node:
            collect(child)

            if child.tail:
                pieces.append(child.tail)

    collect(element)

    return normalize_text("".join(pieces)).strip()


# ---------------------------------------------------------------------------
# Изчислен стил на елемент
# ---------------------------------------------------------------------------

class Style:
    """Свойствата, които документът може да покаже."""

    __slots__ = ("color", "align", "bold", "italic", "underline", "list_type")

    def __init__(self):
        self.color = None
        self.align = None
        self.bold = False
        self.italic = False
        self.underline = False
        self.list_type = "disc"

    def copy(self):
        clone = Style()
        clone.color = self.color
        clone.align = self.align
        clone.bold = self.bold
        clone.italic = self.italic
        clone.underline = self.underline
        clone.list_type = self.list_type

        return clone


def list_type_from(declarations):
    if "list-style-type" in declarations:
        return declarations["list-style-type"].strip().lower()

    if "list-style" in declarations:
        # Съкратеният запис връща типа към началния, ако не е посочен.
        for token in declarations["list-style"].lower().split():
            if token in LIST_TYPES or token == "none":
                return token

        return "disc"

    return None


def compute_style(declarations, parent):
    """Изчислява стила на елемента върху наследеното от родителя."""
    style = parent.copy()

    color = parse_color(declarations.get("color"))

    if color is not None:
        style.color = color

    align = declarations.get("text-align")

    if align:
        style.align = align.strip().lower()

    weight = declarations.get("font-weight")

    if weight:
        weight = weight.strip().lower()

        if weight in ("bold", "bolder"):
            style.bold = True
        elif weight in ("normal", "lighter"):
            style.bold = False
        elif weight.isdigit():
            style.bold = int(weight) >= 600

    font_style = declarations.get("font-style")

    if font_style:
        style.italic = font_style.strip().lower() in ("italic", "oblique")

    decoration = declarations.get("text-decoration") or declarations.get(
        "text-decoration-line"
    )

    if decoration:
        style.underline = "underline" in decoration.lower()

    list_type = list_type_from(declarations)

    if list_type:
        style.list_type = list_type

    flex_align = flex_alignment(declarations)

    if flex_align:
        style.align = flex_align

    return style


FLEX_ALIGNMENTS = {
    "center": "center",
    "flex-end": "right",
    "end": "right",
    "right": "right",
    "flex-start": "left",
    "start": "left",
    "left": "left",
}


def flex_alignment(declarations):
    """
    Превежда подредбата с flexbox в подравняване на текста.

    При ред водещото свойство е justify-content, при колона - align-items.
    """
    display = (declarations.get("display") or "").strip()

    if display not in ("flex", "inline-flex", "grid", "inline-grid"):
        return None

    direction = (declarations.get("flex-direction") or "row").strip()

    if direction.startswith("column"):
        value = (declarations.get("align-items") or "").strip()
    else:
        value = (declarations.get("justify-content") or "").strip()

    return FLEX_ALIGNMENTS.get(value)


# ---------------------------------------------------------------------------
# Изображения
# ---------------------------------------------------------------------------

class ImageLibrary:
    """Декодира вградените Base64 изображения и ги записва без повторения."""

    def __init__(self, output_directory):
        self.output_directory = Path(output_directory)
        self.sources = {}
        self.saved = {}
        self.counter = 0
        self.invalid_count = 0

    def decode_source(self, data_url):
        match = DATA_IMAGE_PATTERN.match((data_url or "").strip())

        if not match:
            return None, None

        encoded = re.sub(r"\s+", "", html.unescape(match.group("data")))

        try:
            raw = base64.b64decode(encoded, validate=True)
        except (binascii.Error, ValueError):
            padding = len(encoded) % 4

            if padding:
                encoded += "=" * (4 - padding)

            try:
                raw = base64.b64decode(encoded)
            except (binascii.Error, ValueError):
                self.invalid_count += 1
                return None, None

        key = hashlib.sha256(raw).hexdigest()

        if key not in self.sources:
            try:
                image = Image.open(io.BytesIO(raw))
                image.load()
                self.sources[key] = image.convert("RGBA")
            except Exception:
                self.invalid_count += 1
                return None, None

        return key, self.sources[key]

    def save(self, image, key, prefix="image"):
        if key in self.saved:
            return self.saved[key]

        self.output_directory.mkdir(parents=True, exist_ok=True)
        self.counter += 1

        path = self.output_directory / "{}_{:03d}.png".format(prefix, self.counter)
        image.save(path)
        self.saved[key] = path

        return path


def crop_image(source_image, declarations, parent_width):
    """
    Пресмята коя част от изображението се вижда и с какъв размер.

    Спазва object-fit: none заедно с object-position, така се държат
    изрязаните от sprite символи.
    """
    source_width, source_height = source_image.size

    box_width = parse_length(declarations.get("width"), parent_width)
    box_height = parse_length(declarations.get("height"))

    fits = (declarations.get("object-fit") or "").strip().lower() == "none"

    image = source_image

    if fits and box_width and box_height:
        position = (declarations.get("object-position") or "0 0").split()

        offset_x = parse_length(position[0]) or 0.0
        offset_y = parse_length(position[1]) if len(position) > 1 else 0.0
        offset_y = offset_y or 0.0

        left = max(0, int(round(-offset_x)))
        top = max(0, int(round(-offset_y)))
        right = min(int(round(left + box_width)), source_width)
        bottom = min(int(round(top + box_height)), source_height)

        if right > left and bottom > top:
            image = source_image.crop((left, top, right, bottom))

        display_width = float(image.width)
        display_height = float(image.height)
    else:
        if box_width and not box_height:
            box_height = source_height * box_width / source_width
        elif box_height and not box_width:
            box_width = source_width * box_height / source_height

        display_width = float(box_width or source_width)
        display_height = float(box_height or source_height)

    zoom = declarations.get("zoom")

    if zoom:
        try:
            factor = float(zoom.strip().rstrip("%")) / (
                100.0 if zoom.strip().endswith("%") else 1.0
            )
            display_width *= factor
            display_height *= factor
        except ValueError:
            pass

    return image, display_width, display_height


def border_color_of(image):
    """
    Връща цвета на плътния фон, ако ръбът на картинката е еднакъв.

    Така се разпознават изрязванията от sprite, който е нарисуван върху
    плътен фон - в браузъра той съвпада с фона на страницата и не се вижда.
    """
    width, height = image.size

    if width < 4 or height < 4:
        return None

    pixels = image.load()

    edge = (
        [(x, 0) for x in range(width)]
        + [(x, height - 1) for x in range(width)]
        + [(0, y) for y in range(height)]
        + [(width - 1, y) for y in range(height)]
    )

    opaque = [pixels[x, y] for x, y in edge if pixels[x, y][3] > 200]

    if len(opaque) < 0.85 * len(edge):
        return None

    reference = opaque[0][:3]

    close = sum(
        1
        for pixel in opaque
        if max(abs(pixel[index] - reference[index]) for index in range(3)) < 20
    )

    if close < 0.85 * len(opaque):
        return None

    return reference


def strip_background(image, backdrop, tolerance=BACKDROP_TOLERANCE):
    """
    Прави фона прозрачен, без да пипа тъмните детайли вътре в рисунката.

    Фонът се разпознава по свързаност с ръба, затова точките по заровете
    и очертанията остават. Меките преходи - сияния и сенки - избледняват
    плавно, защото рисунката е слята върху този фон: разстоянието до него
    дава прозрачността, а цветът се възстановява обратно.
    """
    width, height = image.size
    pixels = image.load()

    distance = Image.new("L", (width, height))
    distance_pixels = distance.load()

    for y in range(height):
        for x in range(width):
            red, green, blue, alpha = pixels[x, y]

            if alpha == 0:
                distance_pixels[x, y] = 255
                continue

            distance_pixels[x, y] = min(
                255,
                max(
                    abs(red - backdrop[0]),
                    abs(green - backdrop[1]),
                    abs(blue - backdrop[2]),
                ),
            )

    near = distance.point(lambda value: 255 if value < tolerance else 0)

    # Рамка от един пиксел, за да тръгне заливането от всички страни.
    mask = Image.new("L", (width + 2, height + 2), 255)
    mask.paste(near, (1, 1))
    ImageDraw.floodfill(mask, (0, 0), 128)
    mask_pixels = mask.load()

    result = image.copy()
    target = result.load()

    for y in range(height):
        for x in range(width):
            if mask_pixels[x + 1, y + 1] != 128:
                continue

            spread = distance_pixels[x, y]

            if spread <= BACKDROP_FLOOR:
                target[x, y] = (0, 0, 0, 0)
                continue

            alpha = min(
                255,
                int(round((spread - BACKDROP_FLOOR) * 255.0 / (tolerance - BACKDROP_FLOOR))),
            )

            if alpha <= 0:
                target[x, y] = (0, 0, 0, 0)
                continue
            red, green, blue, old_alpha = pixels[x, y]
            scale = 255.0 / alpha

            target[x, y] = (
                max(0, min(255, int(backdrop[0] + (red - backdrop[0]) * scale))),
                max(0, min(255, int(backdrop[1] + (green - backdrop[1]) * scale))),
                max(0, min(255, int(backdrop[2] + (blue - backdrop[2]) * scale))),
                min(old_alpha, alpha),
            )

    return result


def sample_pixels(image, size=48):
    sample = image.resize((size, size)).convert("RGBA")
    raw = sample.tobytes()

    return [tuple(raw[start:start + 4]) for start in range(0, len(raw), 4)]


def is_invisible_on_white(image):
    """
    Разпознава иконите, нарисувани за тъмен екран.

    Те са светъл знак върху тъмна полупрозрачна подложка: върху бяло
    подложката избледнява до сиво, а знакът се губи.
    """
    pixels = sample_pixels(image)
    visible = [pixel for pixel in pixels if pixel[3] > 10]

    if not visible:
        return False

    partial = [pixel for pixel in pixels if 10 < pixel[3] < 245]

    # Плътните картинки си носят собствен фон и се виждат.
    if len(visible) > 0.92 * len(pixels):
        return False

    if partial and len(partial) > 0.25 * len(pixels):
        average = sum(luminance(pixel) for pixel in partial) / len(partial)

        if average < 110:
            return True

    return sum(luminance(pixel) for pixel in visible) / len(visible) > 205


def to_dark_glyph(image, color=(26, 26, 26)):
    """Оставя само знака, пребоядисан в тъмно, върху прозрачен фон."""
    grey = image.convert("L")

    visible = grey.point(
        lambda value: 0 if value < 70 else min(255, int((value - 70) * 255 / 130))
    )
    alpha = ImageChops.multiply(visible, image.getchannel("A"))

    glyph = Image.new("RGBA", image.size, color + (255,))
    glyph.putalpha(alpha)

    return glyph


# ---------------------------------------------------------------------------
# Снимане на съставни елементи с браузър
# ---------------------------------------------------------------------------

BROWSER_CANDIDATES = [
    "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
    "/Applications/Chromium.app/Contents/MacOS/Chromium",
    "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
    "/usr/bin/google-chrome",
    "/usr/bin/chromium",
    "/usr/bin/chromium-browser",
]


def find_browser():
    for name in ("google-chrome", "chromium", "chromium-browser"):
        found = shutil.which(name)

        if found:
            return found

    for candidate in BROWSER_CANDIDATES:
        if Path(candidate).exists():
            return candidate

    return None


class ElementRenderer:
    """
    Снима отделен елемент от страницата.

    Нужно е за неща, които Word не може да подреди: текст върху картинка
    и рисунки със SVG. Елементът се снима заедно с веригата от родители,
    за да важат същите CSS правила.
    """

    SCALE = 2

    # Страницата тегли шрифт от мрежата с display=block. Докато шрифтът се
    # чака, текстът е невидим, затова външните @import правила отпадат.
    IMPORT_PATTERN = re.compile(r"@import\s+(?:url\()?[\"'][^\"']+[\"']\)?[^;]*;")

    def __init__(self, stylesheets):
        self.browser = find_browser()
        self.stylesheets = [
            self.IMPORT_PATTERN.sub("", sheet) for sheet in stylesheets
        ]
        self.cache = {}

    @property
    def available(self):
        return self.browser is not None

    # Обвивката пази само id и class, за да важат същите CSS правила.
    # Всичко, което мести или разтяга, се неутрализира - иначе снимката
    # хваща цялата таблица, а не бутона в нея.
    NEUTRAL_STYLE = (
        "display:block;position:static;float:none;width:auto;height:auto;"
        "min-width:0;max-width:none;border:0;padding:0;margin:0;"
        "background:none;box-shadow:none;overflow:visible"
    )

    def wrap_in_context(self, element, fragment):
        """Обгражда елемента с празни копия на родителите му."""
        chain = []
        node = element.getparent()

        while node is not None and isinstance(node.tag, str):
            if node.tag in ("html", "body"):
                break

            chain.append(node)
            node = node.getparent()

        opening = []
        closing = []

        for node in reversed(chain):
            attributes = ""

            for name in ("id", "class"):
                if node.get(name):
                    attributes += ' {}="{}"'.format(name, html.escape(node.get(name)))

            opening.append(
                '<{}{} style="{}">'.format(node.tag, attributes, self.NEUTRAL_STYLE)
            )
            closing.append("</{}>".format(node.tag))

        return "".join(opening) + fragment + "".join(reversed(closing))

    def render(self, element):
        fragment = lxml_html.tostring(element, encoding="unicode")
        key = hashlib.sha256(fragment.encode("utf-8")).hexdigest()

        if key in self.cache:
            return self.cache[key]

        if not self.available:
            self.cache[key] = None
            return None

        styles = "\n".join(
            "<style>{}</style>".format(sheet) for sheet in self.stylesheets
        )
        body = self.wrap_in_context(element, fragment)

        document = (
            "<!doctype html><html><head><meta charset='utf-8'>{}</head>"
            "<body style='margin:0;padding:0;background:transparent'>{}</body></html>"
        ).format(styles, body)

        image = None
        workspace = tempfile.mkdtemp(prefix="html2docx_")

        try:
            page = Path(workspace) / "element.html"
            shot = Path(workspace) / "element.png"
            page.write_text(document, encoding="utf-8")

            subprocess.run(
                [
                    self.browser,
                    "--headless",
                    "--disable-gpu",
                    "--no-sandbox",
                    "--hide-scrollbars",
                    "--force-device-scale-factor={}".format(self.SCALE),
                    "--default-background-color=00000000",
                    "--virtual-time-budget=4000",
                    "--window-size=900,900",
                    "--screenshot={}".format(shot),
                    page.as_uri(),
                ],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=90,
                **hidden_process_flags()
            )

            if shot.exists():
                rendered = Image.open(shot)
                rendered.load()
                rendered = rendered.convert("RGBA")
                bounds = rendered.getbbox()

                if bounds:
                    image = rendered.crop(bounds)
        except (OSError, subprocess.SubprocessError):
            image = None
        finally:
            shutil.rmtree(workspace, ignore_errors=True)

        self.cache[key] = image

        return image


# ---------------------------------------------------------------------------
# Помощни функции за Word
# ---------------------------------------------------------------------------

def set_cell_borders(cell, color, eighths):
    properties = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")

    for edge in ("top", "left", "bottom", "right"):
        line = OxmlElement("w:{}".format(edge))
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), str(eighths))
        line.set(qn("w:space"), "0")
        line.set(qn("w:color"), color)
        borders.append(line)

    properties.append(borders)


def clear_table_borders(table):
    borders = OxmlElement("w:tblBorders")

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        line = OxmlElement("w:{}".format(edge))
        line.set(qn("w:val"), "none")
        line.set(qn("w:sz"), "0")
        borders.append(line)

    table._tbl.tblPr.append(borders)


def set_repeat_header(row):
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def set_paragraph_rule(paragraph, color):
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    line = OxmlElement("w:bottom")
    line.set(qn("w:val"), "single")
    line.set(qn("w:sz"), "6")
    line.set(qn("w:space"), "1")
    line.set(qn("w:color"), color)
    borders.append(line)
    properties.append(borders)


def remove_empty_paragraphs(cell):
    while len(cell.paragraphs) > 1:
        removed = False

        for paragraph in cell.paragraphs:
            if paragraph.text.strip() or paragraph._p.findall(qn("w:r")):
                continue

            paragraph._p.getparent().remove(paragraph._p)
            removed = True
            break

        if not removed:
            break


def set_column_widths(table, widths):
    """
    Задава ширините на колоните.

    При фиксирана подредба Word чете w:tblGrid, а не ширината на всяка
    клетка - затова се задават и двете.
    """
    table.autofit = False

    for index, width in enumerate(widths):
        if index < len(table.columns):
            table.columns[index].width = Inches(width)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths):
                cell.width = Inches(widths[index])


class Container:
    """
    Обвивка около тялото на документа или около клетка от таблица.

    Клетките идват с готов празен параграф - той се използва за първия
    елемент, за да не остават празни редове.
    """

    def __init__(self, part, reuse_first=False):
        self.part = part
        self.pending = part.paragraphs[0] if reuse_first and part.paragraphs else None

    def add_paragraph(self):
        if self.pending is not None:
            paragraph = self.pending
            self.pending = None
            return paragraph

        return self.part.add_paragraph()

    def add_table(self, rows, cols):
        self.pending = None
        return self.part.add_table(rows, cols)


# ---------------------------------------------------------------------------
# Конвертор
# ---------------------------------------------------------------------------

class HtmlToDocxConverter:
    def __init__(self, html_file, images_directory, log=None):
        self.html_file = Path(html_file)
        self.images_directory = Path(images_directory)
        self.log = log or (lambda message: None)

        self.library = ImageLibrary(self.images_directory)
        self.styles = Stylesheet()
        self.renderer = None
        self.document = None
        self.processed = {}
        self.image_count = 0

    # -- четене ----------------------------------------------------------

    def read_html(self):
        for encoding in ("utf-8-sig", "utf-8", "windows-1251", "windows-1252", "latin-1"):
            try:
                return self.html_file.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue

        return self.html_file.read_text(encoding="utf-8", errors="replace")

    def prepare_document(self):
        document = Document()

        section = document.sections[0]
        section.page_width = Mm(PAGE_WIDTH_MM)
        section.page_height = Mm(PAGE_HEIGHT_MM)
        section.top_margin = Inches(PAGE_MARGIN_INCHES)
        section.bottom_margin = Inches(PAGE_MARGIN_INCHES)
        section.left_margin = Inches(PAGE_MARGIN_INCHES)
        section.right_margin = Inches(PAGE_MARGIN_INCHES)

        normal = document.styles["Normal"]
        normal.font.name = BASE_FONT
        normal.font.size = Pt(FONT_SIZE_PT)
        normal.paragraph_format.space_after = Pt(6)

        fonts = normal.element.rPr.rFonts
        fonts.set(qn("w:eastAsia"), BASE_FONT)
        fonts.set(qn("w:cs"), BASE_FONT)

        return document

    def convert(self, output_file):
        source = self.read_html()

        # Без huge_tree lxml спира да чете при няколко мегабайта и
        # мълчаливо връща само началото на страницата.
        parser = lxml_html.HTMLParser(huge_tree=True, recover=True)
        tree = lxml_html.document_fromstring(source, parser=parser)

        stylesheets = [
            node.text_content() for node in tree.iter("style") if node.text_content()
        ]

        for sheet in stylesheets:
            self.styles.add(sheet)

        self.renderer = ElementRenderer(stylesheets)

        if not self.renderer.available:
            self.log(
                "  • Chrome не е намерен - бутоните с текст върху картинка "
                "влизат като картинка плюс текст."
            )

        self.document = self.prepare_document()

        title = tree.find(".//title")

        if title is not None and title.text and title.text.strip():
            self.add_title(title.text.strip())

        body = tree.find("body")
        root = body if body is not None else tree

        base = compute_style(self.declarations(root), Style())
        base.color = readable_on_white(base.color)

        for child in root:
            self.walk(child, Container(self.document), base)

        if output_file is not None:
            self.document.save(str(output_file))

        return output_file

    def add_title(self, text):
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(12)

        run = paragraph.add_run(text)
        run.bold = True
        self.apply_font(run)

    # -- достъп до стиловете ---------------------------------------------

    def declarations(self, element):
        return self.styles.declarations_for(element)

    def is_hidden(self, element):
        declarations = self.declarations(element)

        if (declarations.get("display") or "").strip().lower() == "none":
            return True

        return (declarations.get("visibility") or "").strip().lower() == "hidden"

    def is_inline(self, element):
        display = (self.declarations(element).get("display") or "").strip().lower()

        return display.startswith("inline")

    def parent_width(self, element):
        parent = element.getparent()

        if parent is None:
            return None

        return parse_length(self.declarations(parent).get("width"))

    # -- разпознаване на блокове -----------------------------------------

    def has_block_content(self, element):
        for node in element.iterdescendants():
            if not isinstance(node.tag, str):
                continue

            if node.tag in BLOCK_TAGS:
                return True

            if node.tag == "img" and not self.is_inline(node):
                return True

        return False

    def is_rule(self, element):
        """Тънка декоративна ивица без съдържание - разделителна линия."""
        if visible_text(element):
            return False

        declarations = self.declarations(element)
        height = parse_length(declarations.get("height"))

        if height is None or height > 8:
            return False

        return bool(
            declarations.get("background-image") or declarations.get("background-color")
        )

    def is_stacked(self, element):
        """
        Картинка с текст върху нея - подрежда се само от браузър.

        Елементът трябва сам да е позициониран: тогава той е опората на
        абсолютно наместените си деца. Иначе всеки родител нагоре по
        дървото би минал за такъв и цялата страница би станала снимка.
        """
        position = (self.declarations(element).get("position") or "").strip().lower()

        if position not in ("relative", "absolute", "fixed", "sticky"):
            return False

        for node in element.iterdescendants():
            if not isinstance(node.tag, str):
                continue

            child_position = (
                self.declarations(node).get("position") or ""
            ).strip().lower()

            if child_position in ("absolute", "fixed"):
                return True

        return False

    def flex_columns(self, element):
        """Брой колони при подредба с flex-wrap и flex-basis."""
        declarations = self.declarations(element)

        if (declarations.get("display") or "").strip() != "flex":
            return None

        if (declarations.get("flex-wrap") or "").strip() != "wrap":
            return None

        for child in element:
            if not isinstance(child.tag, str):
                continue

            basis = (self.declarations(child).get("flex-basis") or "").strip()
            match = re.fullmatch(r"(\d+(?:\.\d+)?)%", basis)

            if match:
                share = float(match.group(1))

                if share > 0:
                    return max(1, int(round(100.0 / share)))

        return None

    def flex_row(self, element):
        """Колоните на подредба в ред, ако има поне две със съдържание."""
        declarations = self.declarations(element)

        if (declarations.get("display") or "").strip() != "flex":
            return None

        direction = (declarations.get("flex-direction") or "row").strip()

        if direction != "row" or (declarations.get("flex-wrap") or "").strip() == "wrap":
            return None

        children = [
            child
            for child in element
            if isinstance(child.tag, str)
            and not self.is_hidden(child)
            and self.has_content(child)
        ]

        if len(children) < 2:
            return None

        return children

    def has_content(self, element):
        if visible_text(element):
            return True

        return any(True for _ in element.iter("img")) or any(
            True for _ in element.iter("svg")
        )

    # -- обхождане -------------------------------------------------------

    def walk(self, element, container, style):
        if not isinstance(element.tag, str):
            return

        tag = element.tag.lower()

        if tag in SKIP_TAGS or self.is_hidden(element):
            return

        own = compute_style(self.declarations(element), style)

        if tag == "table":
            self.emit_table(element, container, own)
            return

        if tag == "svg":
            self.emit_rendered(element, container, own)
            return

        if tag in HEADING_TAGS:
            self.emit_paragraph(element, container, own, heading=True)
            return

        if tag == "p":
            self.emit_paragraph(element, container, own)
            return

        if tag == "li":
            self.emit_paragraph(element, container, own, bullet=True)
            return

        if tag == "img":
            self.emit_image(element, container, own)
            return

        if tag == "hr":
            self.emit_rule(container)
            return

        if tag == "br":
            return

        if tag in ("ul", "ol"):
            self.walk_children(element, container, own)
            return

        if self.is_rule(element):
            self.emit_rule(container)
            return

        if self.is_stacked(element):
            self.emit_rendered(element, container, own)
            return

        if not self.has_block_content(element):
            self.emit_paragraph(element, container, own)
            return

        columns = self.flex_columns(element)

        if columns:
            self.emit_grid(element, container, own, columns)
            return

        row = self.flex_row(element)

        if row:
            self.emit_columns(row, container, own)
            return

        self.walk_children(element, container, own)

    def walk_children(self, element, container, style):
        if normalize_text(element.text).strip():
            self.emit_text(element.text, container, style)

        for child in element:
            self.walk(child, container, style)

            if isinstance(child.tag, str) and normalize_text(child.tail).strip():
                self.emit_text(child.tail, container, style)

    # -- блокове ---------------------------------------------------------

    def new_paragraph(self, container, style, declarations=None):
        paragraph = container.add_paragraph()
        paragraph.alignment = self.alignment_of(style, declarations)

        return paragraph

    def alignment_of(self, style, declarations=None):
        """
        Подравняването идва от text-align, но блоковете с margin: auto
        се центрират - точно както прави браузърът с картинките.
        """
        if declarations:
            sides = self.auto_margins(declarations)

            if sides == (True, True):
                return WD_ALIGN_PARAGRAPH.CENTER

            if sides == (True, False):
                return WD_ALIGN_PARAGRAPH.RIGHT

            if sides == (False, True):
                return WD_ALIGN_PARAGRAPH.LEFT

        return paragraph_alignment(style.align or "left")

    def auto_margins(self, declarations):
        """Връща (лявото е auto, дясното е auto)."""
        shorthand = declarations.get("margin-inline") or declarations.get("margin")

        if shorthand:
            parts = shorthand.split()

            if all(part == "auto" for part in parts):
                return True, True

            if len(parts) >= 2 and parts[1] == "auto":
                return True, True

        left = (declarations.get("margin-left") or "").strip() == "auto"
        right = (declarations.get("margin-right") or "").strip() == "auto"

        return left, right

    def apply_font(self, run, style=None):
        run.font.name = BASE_FONT
        run.font.size = Pt(FONT_SIZE_PT)

        if style is not None:
            run.bold = style.bold
            run.italic = style.italic
            run.underline = style.underline
            run.font.color.rgb = RGBColor(*readable_on_white(style.color))

        return run

    def emit_text(self, text, container, style):
        paragraph = self.new_paragraph(container, style)
        self.apply_font(paragraph.add_run(normalize_text(text).strip()), style)

    def emit_paragraph(self, element, container, style, heading=False, bullet=False):
        paragraph = self.new_paragraph(container, style)

        if heading:
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.keep_with_next = True

        if bullet and style.list_type != "none":
            marker = LIST_TYPES.get(style.list_type, "•")
            colors = self.styles.declarations_for(element, pseudo="marker")
            color = parse_color(colors.get("color")) or style.color

            run = paragraph.add_run(marker + "  ")
            run.font.name = BASE_FONT
            run.font.size = Pt(FONT_SIZE_PT)
            run.font.color.rgb = RGBColor(*readable_on_white(color))

        if not self.write_inline(element, paragraph, style) and not bullet:
            self.drop(paragraph)

    def emit_rule(self, container):
        paragraph = container.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(8)
        set_paragraph_rule(paragraph, "BFBFBF")

    def emit_image(self, element, container, style, width_limit=None):
        prepared = self.prepare_image(element)

        if prepared is None:
            return

        path, width_px, _ = prepared
        width = min(width_px * INCHES_PER_CSS_PX, width_limit or CONTENT_WIDTH_INCHES)

        paragraph = self.new_paragraph(container, style, self.declarations(element))
        paragraph.paragraph_format.space_after = Pt(4)

        try:
            paragraph.add_run().add_picture(str(path), width=Inches(width))
            self.image_count += 1
        except Exception:
            self.drop(paragraph)

    def emit_rendered(self, element, container, style, width_limit=None):
        rendered = self.renderer.render(element)

        if rendered is None:
            self.emit_rendered_fallback(element, container, style)
            return

        key = hashlib.sha256(rendered.tobytes()).hexdigest()
        path = self.library.save(rendered, key, prefix="control")

        width_px = rendered.width / ElementRenderer.SCALE
        width = min(width_px * INCHES_PER_CSS_PX, width_limit or CONTENT_WIDTH_INCHES)

        paragraph = self.new_paragraph(container, style, self.declarations(element))
        paragraph.paragraph_format.space_after = Pt(4)

        try:
            paragraph.add_run().add_picture(str(path), width=Inches(width))
            self.image_count += 1
        except Exception:
            self.drop(paragraph)

    def emit_rendered_fallback(self, element, container, style):
        for image in element.iter("img"):
            self.emit_image(image, container, style)

        text = visible_text(element)

        if text:
            paragraph = self.new_paragraph(container, style)
            self.apply_font(paragraph.add_run(text), style)

    def emit_grid(self, element, container, style, columns):
        children = [
            child
            for child in element
            if isinstance(child.tag, str) and not self.is_hidden(child)
        ]

        if not children:
            return

        rows = (len(children) + columns - 1) // columns
        table = container.add_table(rows=rows, cols=columns)
        clear_table_borders(table)

        for index, child in enumerate(children):
            cell = table.cell(index // columns, index % columns)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            self.walk(child, Container(cell, reuse_first=True), style)
            remove_empty_paragraphs(cell)

        set_column_widths(table, [CONTENT_WIDTH_INCHES / columns] * columns)

    def emit_columns(self, children, container, style):
        table = container.add_table(rows=1, cols=len(children))
        clear_table_borders(table)

        for index, child in enumerate(children):
            cell = table.cell(0, index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            self.walk(child, Container(cell, reuse_first=True), style)
            remove_empty_paragraphs(cell)

        set_column_widths(
            table,
            [CONTENT_WIDTH_INCHES / len(children)] * len(children),
        )

    # -- таблици ---------------------------------------------------------

    def emit_table(self, element, container, style):
        rows = list(element.iter("tr"))

        if not rows:
            return

        cells_per_row = [
            [cell for cell in row if cell.tag in ("td", "th")] for row in rows
        ]
        column_count = max((len(cells) for cells in cells_per_row), default=0)

        if column_count == 0:
            return

        widths = self.column_widths(cells_per_row, column_count)

        table = container.add_table(rows=len(rows), cols=column_count)

        for row_index, cells in enumerate(cells_per_row):
            for column_index in range(column_count):
                cell = table.cell(row_index, column_index)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                if column_index >= len(cells):
                    continue

                source = cells[column_index]
                cell_style = compute_style(self.declarations(source), style)

                self.apply_cell_border(cell, source, cell_style)
                self.fill_cell(source, cell, cell_style, widths[column_index])
                remove_empty_paragraphs(cell)

            if row_index == 0:
                set_repeat_header(table.rows[0])

        set_column_widths(table, widths)

    def apply_cell_border(self, cell, source, style):
        declarations = self.declarations(source)

        width = parse_length(
            declarations.get("border-width") or declarations.get("border")
        )

        if not width:
            return

        color = parse_color(declarations.get("border-color"))

        # Без изричен цвят рамката ползва цвета на текста - както в браузъра.
        if color is None:
            color = style.color

        eighths = max(2, int(round(width * 0.75 * 8)))
        set_cell_borders(cell, to_hex(readable_on_white(color)), eighths)

    def column_widths(self, cells_per_row, column_count):
        """
        Пресмята ширините по съдържанието, както прави браузърът.

        Картинките дават естествената си ширина, текстът - приблизителна
        по броя знаци, с таван, за да не изяде реда една дълга клетка.
        """
        natural = [1.0] * column_count

        for cells in cells_per_row:
            for index, cell in enumerate(cells[:column_count]):
                widest = 0.0

                # Картинките в клетка стоят една под друга, затова
                # решаваща е най-широката, а не сборът им.
                for image in cell.iter("img"):
                    declarations = self.declarations(image)
                    widest = max(widest, parse_length(declarations.get("width")) or 60.0)

                text = len(visible_text(cell)) * 7.0
                natural[index] = max(
                    natural[index], widest + 20.0, min(text, TEXT_WIDTH_CAP)
                )

        total = sum(natural)

        return [CONTENT_WIDTH_INCHES * value / total for value in natural]

    def fill_cell(self, source, cell, style, width_inches):
        target = Container(cell, reuse_first=True)
        limit = max(0.3, width_inches - 0.1)

        if not self.has_block_content(source) and not self.is_stacked(source):
            paragraph = self.new_paragraph(target, style)

            if not self.write_inline(source, paragraph, style):
                self.drop(paragraph)

            return

        if normalize_text(source.text).strip():
            self.emit_text(source.text, target, style)

        for child in source:
            self.walk_in_cell(child, target, style, limit)

            if isinstance(child.tag, str) and normalize_text(child.tail).strip():
                self.emit_text(child.tail, target, style)

    def walk_in_cell(self, element, target, style, limit):
        if not isinstance(element.tag, str):
            return

        tag = element.tag.lower()

        if tag in SKIP_TAGS or self.is_hidden(element):
            return

        own = compute_style(self.declarations(element), style)

        if tag == "svg":
            self.emit_rendered(element, target, own, width_limit=limit)
            return

        if tag == "img":
            self.emit_image(element, target, own, width_limit=limit)
            return

        if tag == "br":
            return

        if self.is_stacked(element):
            self.emit_rendered(element, target, own, width_limit=limit)
            return

        if not self.has_block_content(element):
            paragraph = self.new_paragraph(target, own)

            if not self.write_inline(element, paragraph, own):
                self.drop(paragraph)

            return

        for child in element:
            self.walk_in_cell(child, target, own, limit)

    # -- текст в един параграф -------------------------------------------

    def write_inline(self, element, paragraph, style):
        state = {"written": False, "space": True}
        self.collect_inline(element, paragraph, style, state)

        return state["written"]

    def collect_inline(self, element, paragraph, style, state):
        own = compute_style(self.declarations(element), style)

        self.append_text(paragraph, element.text, own, state)

        for child in element:
            if not isinstance(child.tag, str):
                self.append_text(paragraph, child.tail, own, state)
                continue

            tag = child.tag.lower()

            if tag in SKIP_TAGS or self.is_hidden(child):
                pass
            elif tag == "br":
                paragraph.add_run().add_break()
                state["space"] = True
            elif tag == "img":
                self.append_image(paragraph, child, state)
            elif tag != "svg":
                self.collect_inline(child, paragraph, own, state)

            self.append_text(paragraph, child.tail, own, state)

    def append_text(self, paragraph, text, style, state):
        text = normalize_text(text)

        if state["space"]:
            text = text.lstrip()

        if not text:
            return

        self.apply_font(paragraph.add_run(text), style)

        state["written"] = True
        state["space"] = text.endswith(" ")

    def append_image(self, paragraph, element, state):
        prepared = self.prepare_image(element)

        if prepared is None:
            return

        path, width_px, height_px = prepared

        width = width_px * INCHES_PER_CSS_PX
        height = height_px * INCHES_PER_CSS_PX

        if width > CONTENT_WIDTH_INCHES:
            height *= CONTENT_WIDTH_INCHES / width
            width = CONTENT_WIDTH_INCHES

        try:
            paragraph.add_run().add_picture(str(path), height=Inches(height))
            self.image_count += 1
            state["written"] = True
            state["space"] = False
        except Exception:
            pass

    # -- една картинка ---------------------------------------------------

    def prepare_image(self, element):
        source_url = element.get("src") or ""

        if not source_url.startswith("data:image"):
            return None

        key, source_image = self.library.decode_source(source_url)

        if source_image is None:
            return None

        declarations = self.declarations(element)

        image, width_px, height_px = crop_image(
            source_image,
            declarations,
            self.parent_width(element),
        )

        if width_px <= 0 or height_px <= 0:
            return None

        crop_key = "{}:{}x{}:{}".format(
            key,
            image.width,
            image.height,
            hashlib.sha256(image.tobytes()).hexdigest()[:16],
        )

        # Едно изрязване се среща десетки пъти в страницата, а обработката
        # му е скъпа - затова се смята само веднъж.
        if crop_key in self.processed:
            return self.processed[crop_key], width_px, height_px

        backdrop = border_color_of(image)

        if backdrop is not None:
            image = strip_background(image, backdrop)
        elif FIX_INVISIBLE_IMAGES and is_invisible_on_white(image):
            image = to_dark_glyph(image)

        path = self.library.save(image, crop_key)
        self.processed[crop_key] = path

        return path, width_px, height_px

    def drop(self, paragraph):
        parent = paragraph._p.getparent()

        if parent is not None:
            parent.remove(paragraph._p)


# ---------------------------------------------------------------------------
# Публичен вход
# ---------------------------------------------------------------------------

def create_output_path(html_file, suffix):
    """
    report.html -> report.docx

    Ако файлът вече съществува: report_2.docx, report_3.docx ...
    """
    html_file = Path(html_file)
    candidate = html_file.parent / "{}{}".format(html_file.stem, suffix)

    if not candidate.exists():
        return candidate

    counter = 2

    while True:
        candidate = html_file.parent / "{}_{}{}".format(html_file.stem, counter, suffix)

        if not candidate.exists():
            return candidate

        counter += 1


def convert_file(
    html_file,
    output_file=None,
    images_directory=None,
    write_document=True,
    log=None,
):
    """
    Преобразува един HTML файл в .docx и записва снимките му като PNG.

    write_document=False минава по същия път, но записва само снимките.

    Връща (път до документа или None, папка със снимки, брой снимки).
    """
    html_file = Path(html_file)

    if write_document and output_file is None:
        output_file = create_output_path(html_file, ".docx")

    if not write_document:
        output_file = None

    if images_directory is None:
        images_directory = create_output_path(html_file, "_images")

    converter = HtmlToDocxConverter(html_file, images_directory, log=log)
    converter.convert(output_file)

    return output_file, Path(images_directory), converter.library.counter


# =========================================================================
# Приложение
# =========================================================================

APP_TITLE = "HTML Image Extractor"


# Намира изображения от типа:
# data:image/png;base64,iVBORw0KGgo...
#
# Работи както в <img src="...">, така и в CSS background-image.
DATA_IMAGE_SCAN_PATTERN = re.compile(
    r"data:image/"
    r"(?P<mime>[a-zA-Z0-9.+_-]+)"
    r"(?:;[^,]*)?"
    r";base64,"
    r"(?P<data>[a-zA-Z0-9+/=\s]+)",
    re.IGNORECASE,
)


EXTENSIONS = {
    "png": "png",
    "jpeg": "jpg",
    "jpg": "jpg",
    "gif": "gif",
    "webp": "webp",
    "bmp": "bmp",
    "svg+xml": "svg",
    "x-icon": "ico",
    "vnd.microsoft.icon": "ico",
    "tiff": "tiff",
    "avif": "avif",
}


def read_html_file(file_path):
    # Опитва няколко често срещани кодировки.
    encodings = [
        "utf-8-sig",
        "utf-8",
        "windows-1251",
        "windows-1252",
        "latin-1",
    ]

    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as file:
                return file.read()
        except UnicodeDecodeError:
            continue

    with open(file_path, "r", encoding="utf-8", errors="replace") as file:
        return file.read()


def normalize_base64(encoded_data):
    encoded_data = html.unescape(encoded_data)
    return re.sub(r"\s+", "", encoded_data)


def get_extension(mime_subtype):
    mime_subtype = mime_subtype.lower().strip()
    return EXTENSIONS.get(mime_subtype, "bin")


def create_output_directory(html_file):
    """
    report.html -> report_images

    Ако папката вече съществува:
    report_images_2
    report_images_3
    """
    base_directory = html_file.parent / "{}_images".format(html_file.stem)

    if not base_directory.exists():
        base_directory.mkdir(parents=True)
        return base_directory

    counter = 2

    while True:
        directory = html_file.parent / "{}_images_{}".format(
            html_file.stem,
            counter,
        )

        if not directory.exists():
            directory.mkdir(parents=True)
            return directory

        counter += 1


def decode_image(encoded_data):
    """
    Декодира Base64 съдържанието.

    Първо опитва строг режим, след това добавя липсващото
    Base64 padding, ако е необходимо.
    """
    encoded_data = normalize_base64(encoded_data)

    try:
        return base64.b64decode(encoded_data, validate=True)
    except (binascii.Error, ValueError):
        pass

    missing_padding = len(encoded_data) % 4

    if missing_padding:
        encoded_data += "=" * (4 - missing_padding)

    try:
        return base64.b64decode(encoded_data)
    except (binascii.Error, ValueError):
        return None


def extract_images(html_file):
    """
    Извлича Base64 изображенията от един HTML файл.

    Връща:
    saved_count, duplicate_count, invalid_count, output_directory
    """
    html_content = read_html_file(html_file)
    html_content = html.unescape(html_content)

    matches = list(DATA_IMAGE_SCAN_PATTERN.finditer(html_content))

    if not matches:
        return 0, 0, 0, None

    output_directory = create_output_directory(html_file)

    saved_count = 0
    duplicate_count = 0
    invalid_count = 0
    image_hashes = set()

    for match in matches:
        mime_subtype = match.group("mime")
        encoded_data = match.group("data")

        image_bytes = decode_image(encoded_data)

        if not image_bytes:
            invalid_count += 1
            continue

        image_hash = hashlib.sha256(image_bytes).hexdigest()

        if image_hash in image_hashes:
            duplicate_count += 1
            continue

        image_hashes.add(image_hash)

        extension = get_extension(mime_subtype)
        saved_count += 1

        output_file = output_directory / "image_{:03d}.{}".format(
            saved_count,
            extension,
        )

        with open(output_file, "wb") as file:
            file.write(image_bytes)

    if saved_count == 0:
        try:
            output_directory.rmdir()
        except OSError:
            pass

        return 0, duplicate_count, invalid_count, None

    return (
        saved_count,
        duplicate_count,
        invalid_count,
        output_directory,
    )


def open_directory(directory):
    try:
        if sys.platform.startswith("win"):
            os.startfile(str(directory))
        elif sys.platform == "darwin":
            subprocess.Popen(["open", str(directory)])
        else:
            subprocess.Popen(["xdg-open", str(directory)])
    except Exception as error:
        messagebox.showerror(
            APP_TITLE,
            "Папката не може да бъде отворена:\n{}".format(error),
        )


class HTMLImageExtractorApp:
    def __init__(self, root):
        self.root = root
        self.last_output_directory = None

        self.save_images = tk.BooleanVar(value=True)
        self.save_document = tk.BooleanVar(value=CONVERTER_ERROR is None)
        self.messages = queue.Queue()

        self.root.title(APP_TITLE)
        self.root.geometry("680x580")
        self.root.minsize(580, 500)

        self.build_interface()

    def build_interface(self):
        main_frame = ttk.Frame(self.root, padding=24)
        main_frame.pack(fill="both", expand=True)

        title_label = ttk.Label(
            main_frame,
            text="HTML към снимки и Word",
            font=("Segoe UI", 18, "bold"),
        )
        title_label.pack(pady=(0, 8))

        description_label = ttk.Label(
            main_frame,
            text=(
                "Избери един или повече HTML файлове.\n"
                "Вградените Base64 снимки се записват в нова папка, "
                "а текстът и таблиците - в документ на Word."
            ),
            justify="center",
        )
        description_label.pack(pady=(0, 14))

        self.build_options(main_frame)

        self.select_button = ttk.Button(
            main_frame,
            text="Избери HTML файлове",
            command=self.select_files,
        )
        self.select_button.pack(ipadx=24, ipady=10, pady=(0, 16))

        self.progress = ttk.Progressbar(
            main_frame,
            mode="determinate",
        )
        self.progress.pack(fill="x", pady=(0, 14))

        log_frame = ttk.LabelFrame(
            main_frame,
            text="Резултат",
            padding=10,
        )
        log_frame.pack(fill="both", expand=True)

        self.log_text = tk.Text(
            log_frame,
            height=12,
            wrap="word",
            state="disabled",
        )
        self.log_text.pack(side="left", fill="both", expand=True)

        scrollbar = ttk.Scrollbar(
            log_frame,
            orient="vertical",
            command=self.log_text.yview,
        )
        scrollbar.pack(side="right", fill="y")

        self.log_text.configure(yscrollcommand=scrollbar.set)

        self.open_folder_button = ttk.Button(
            main_frame,
            text="Отвори последната папка",
            command=self.open_last_directory,
            state="disabled",
        )
        self.open_folder_button.pack(pady=(14, 0))

    def build_options(self, parent):
        options_frame = ttk.LabelFrame(parent, text="Какво да се създаде", padding=10)
        options_frame.pack(fill="x", pady=(0, 16))

        ttk.Checkbutton(
            options_frame,
            text="Снимки (PNG) - изрязани така, както се виждат на страницата",
            variable=self.save_images,
        ).pack(anchor="w")

        document_checkbox = ttk.Checkbutton(
            options_frame,
            text="Документ на Word (.docx) - текст, таблици и снимки",
            variable=self.save_document,
        )
        document_checkbox.pack(anchor="w", pady=(4, 0))

        if CONVERTER_ERROR is not None:
            document_checkbox.configure(state="disabled")

            ttk.Label(
                options_frame,
                text=(
                    "Компонентите за Word не можаха да се подготвят:\n"
                    "{}\n\n"
                    "Пусни програмата пак с включен интернет."
                ).format(CONVERTER_ERROR),
                foreground="#a05000",
                justify="left",
                wraplength=560,
            ).pack(anchor="w", pady=(6, 0))

    def add_log(self, message):
        self.log_text.configure(state="normal")
        self.log_text.insert("end", message + "\n")
        self.log_text.see("end")
        self.log_text.configure(state="disabled")

    def clear_log(self):
        self.log_text.configure(state="normal")
        self.log_text.delete("1.0", "end")
        self.log_text.configure(state="disabled")

    def select_files(self):
        selected_files = filedialog.askopenfilenames(
            title="Избери HTML файлове",
            filetypes=[
                ("HTML файлове", "*.html *.htm"),
                ("Всички файлове", "*.*"),
            ],
        )

        if not selected_files:
            return

        files = [Path(file_path) for file_path in selected_files]
        self.process_files(files)

    # -- пускане на обработката ------------------------------------------

    def process_files(self, files):
        want_images = self.save_images.get()
        want_document = self.save_document.get() and CONVERTER_ERROR is None

        if not want_images and not want_document:
            messagebox.showinfo(
                APP_TITLE,
                "Избери поне едно от двете - снимки или документ.",
            )
            return

        self.clear_log()

        self.select_button.configure(state="disabled")
        self.open_folder_button.configure(state="disabled")
        self.last_output_directory = None

        self.progress.configure(maximum=len(files), value=0)

        # Големите файлове се обработват десетки секунди, затова работата
        # върви в отделна нишка, а съобщенията се предават през опашка.
        self.messages = queue.Queue()

        worker = threading.Thread(
            target=self.run_batch,
            args=(files, want_images, want_document),
            daemon=True,
        )
        worker.start()

        self.root.after(100, self.drain_messages)

    def post(self, kind, payload):
        self.messages.put((kind, payload))

    def drain_messages(self):
        """Пренася съобщенията от работната нишка към прозореца."""
        finished = False
        summary = ""

        try:
            while True:
                kind, payload = self.messages.get_nowait()

                if kind == "log":
                    self.add_log(payload)
                elif kind == "progress":
                    self.progress.configure(value=payload)
                elif kind == "directory":
                    self.last_output_directory = payload
                elif kind == "done":
                    finished = True
                    summary = payload
        except queue.Empty:
            pass

        if not finished:
            self.root.after(100, self.drain_messages)
            return

        self.select_button.configure(state="normal")

        if self.last_output_directory is not None:
            self.open_folder_button.configure(state="normal")

        messagebox.showinfo(APP_TITLE, summary)

    # -- самата обработка (върви в отделна нишка) ------------------------

    def run_batch(self, files, want_images, want_document):
        total_images = 0
        total_documents = 0
        successful_files = 0
        failed_files = 0

        for index, html_file in enumerate(files, start=1):
            self.post("log", "Обработва се: {}".format(html_file.name))

            try:
                images, document = self.process_one(
                    html_file,
                    want_images,
                    want_document,
                )

                total_images += images
                total_documents += 1 if document else 0

                if images or document:
                    successful_files += 1

            except PermissionError:
                failed_files += 1
                self.post("log", "  ✗ Няма разрешение за запис в тази папка.")

            except OSError as error:
                failed_files += 1
                self.post("log", "  ✗ Файлова грешка: {}".format(error))

            except Exception as error:
                failed_files += 1
                self.post("log", "  ✗ Неочаквана грешка: {}".format(error))

            self.post("log", "")
            self.post("progress", index)

        summary = (
            "Обработката приключи.\n\n"
            "Записани снимки: {}\n"
            "Създадени документи: {}\n"
            "Успешни HTML файлове: {}\n"
            "Файлове с грешка: {}"
        ).format(
            total_images,
            total_documents,
            successful_files,
            failed_files,
        )

        self.post("done", summary)

    def process_one(self, html_file, want_images, want_document):
        """Обработва един файл и връща (брой снимки, път до документа)."""
        if CONVERTER_ERROR is not None:
            return self.process_without_converter(html_file)

        document, images_directory, image_count = convert_file(
            html_file,
            write_document=want_document,
            log=lambda message: self.post("log", message),
        )

        if not want_images and images_directory.exists():
            shutil.rmtree(images_directory, ignore_errors=True)
        elif image_count:
            self.post("log", "  ✓ Записани снимки: {}".format(image_count))
            self.post("log", "  • Папка: {}".format(images_directory))
            self.post("directory", images_directory)
        else:
            self.post("log", "  Няма намерени валидни Base64 снимки.")

        if document is not None:
            self.post("log", "  ✓ Документ: {}".format(document.name))
            self.post("directory", document.parent)

        return (image_count if want_images else 0), document

    def process_without_converter(self, html_file):
        """Резервен път: само изрязване на Base64 снимките."""
        saved_count, duplicate_count, invalid_count, directory = extract_images(
            html_file
        )

        if saved_count:
            self.post("log", "  ✓ Записани снимки: {}".format(saved_count))

            if duplicate_count:
                self.post(
                    "log",
                    "  • Пропуснати дубликати: {}".format(duplicate_count),
                )

            if invalid_count:
                self.post(
                    "log",
                    "  • Невалидни изображения: {}".format(invalid_count),
                )

            self.post("log", "  • Папка: {}".format(directory))
            self.post("directory", directory)
        else:
            self.post("log", "  Няма намерени валидни Base64 снимки.")

        return saved_count, None

    def open_last_directory(self):
        if self.last_output_directory is not None:
            open_directory(self.last_output_directory)

def main():
    # С подадени файлове работи от командния ред, иначе отваря прозореца.
    if len(sys.argv) > 1:
        if CONVERTER_ERROR is not None:
            print("Компонентите за Word не са готови: {}".format(CONVERTER_ERROR))
            print("Пусни командата пак с включен интернет.")
            return 1

        for name in sys.argv[1:]:
            document, images, count = convert_file(name, log=print)
            print("{} -> {} ({} снимки в {})".format(name, document, count, images))

        return 0

    root = tk.Tk()

    try:
        if sys.platform.startswith("win"):
            ttk.Style().theme_use("vista")
    except tk.TclError:
        pass

    HTMLImageExtractorApp(root)
    root.mainloop()

    return 0


if __name__ == "__main__":
    sys.exit(main())
